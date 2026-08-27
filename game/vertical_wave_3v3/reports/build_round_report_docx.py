from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[3]
GAME_DIR = ROOT / "game" / "vertical_wave_3v3"
FIG_DIR = GAME_DIR / "figures"
OUT_PATH = GAME_DIR / "reports" / "vertical_wave_3v3_round_report.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        rest = text[len(bold_prefix):]
        paragraph.add_run(rest)
    else:
        paragraph.add_run(text)
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_picture_if_exists(doc: Document, path: Path, caption: str, width: float = 6.2) -> None:
    if not path.exists():
        add_para(doc, f"缺失图片：{path}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr_cells[idx], header, bold=True)
        set_cell_shading(hdr_cells[idx], "D9EAF7")
        if widths:
            hdr_cells[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            if widths:
                cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()


def load_rounds() -> list[dict]:
    with (GAME_DIR / "round_history.json").open("r", encoding="utf-8") as handle:
        return json.load(handle)


def build_doc() -> None:
    rounds = load_rounds()
    executed = [row for row in rounds if row.get("status") in {"PASS", "FAIL_STOP"}]

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("vertical_wave_3v3 轮次迭代科研报告")
    title_run.bold = True
    title_run.font.name = "Microsoft YaHei"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    title_run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("输入设定、策略迭代、可视化结果与科研阐述")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    add_para(doc, "结论摘要：本轮执行在 Round 4 续跑后蓝方达标，但 Round 5 红方在 10 次受控策略代码迭代后仍未反超蓝方，因此最终状态为 FAIL_STOP。该结果支持“任务存在有效耦合对抗，蓝方 escort/guard 策略在当前有限搜索空间中更稳，但当前策略组合尚不满足经验近似纳什稳定”的科研阐述。")

    add_heading(doc, "1. 输入任务与场景设定", 1)
    add_para(doc, "用户输入：红方 3 机（2 赛车机 + 1 防守机）对阵蓝方 3 机（2 赛车机 + 1 防守机），vertical_wave 布局，随机出生，超时 600 步。需要多车道分配和前视碰撞检测。")
    add_bullet(doc, "形式化类型：POSG，多智能体对抗博弈。")
    add_bullet(doc, "核心目标：红蓝双方交替做固定对手下的策略改进，比较 red_utility 与 blue_utility。")
    add_bullet(doc, "硬约束：collision_rate <= 0.05，out_of_bounds_rate <= 0.01，action_violation_rate == 0。")
    add_bullet(doc, "博弈存在性门禁：目标侧空场耦合 Delta = U_side(side, empty_opponent) - U_side(side, true_opponent) 必须大于 0。")

    add_heading(doc, "2. 方法概述", 1)
    add_para(doc, "本实验没有使用神经网络强化学习，而采用规则安全控制器 + 交替最佳响应 + 参数 sweep + 受控策略代码迭代。")
    add_bullet(doc, "基础控制器：SafeRulePolicy，提供多车道分配、前视碰撞检测、边界保护、制动 fallback 和动作裁剪。")
    add_bullet(doc, "红蓝显式拆分：每轮策略包包含 RedPolicy、BluePolicy 和 PolicyClass。PolicyClass 只做分发、fallback 与 action clipping。")
    add_bullet(doc, "交替最佳响应：每轮只优化当前目标侧，另一侧冻结为上一轮 best policy/config。")
    add_bullet(doc, "参数 sweep：先枚举目标侧速度、风险边界、车道间距、防守模式等参数。")
    add_bullet(doc, "受控代码迭代：参数 sweep 不达标时，仅允许修改当前目标侧策略逻辑；代码迭代上限最终设为 10。")

    add_heading(doc, "3. 逐轮结果总览", 1)
    rows = []
    for row in executed:
        rows.append([
            str(row["round_id"]),
            str(row["target_side"]),
            str(row["status"]),
            f'{row["red_utility"]:.3f}',
            f'{row["blue_utility"]:.3f}',
            f'{row["target_margin"]:.3f}',
            f'{row["coupling_delta"]:.3f}',
            str(int(row["policy_code_edits"])),
        ])
    add_table(
        doc,
        ["Round", "目标侧", "状态", "红方效用", "蓝方效用", "目标 margin", "耦合 Delta", "代码迭代"],
        rows,
        [0.65, 0.8, 1.0, 0.9, 0.9, 1.0, 1.0, 0.85],
    )

    add_heading(doc, "4. 每轮策略改进内容", 1)
    add_heading(doc, "Round 1：初始基线", 2)
    add_para(doc, "建立 vertical_wave_3v3 的初始红蓝组合，用于判断初始劣势方和后续交替优化顺序。本轮在本次续跑前已完成，因此报告中作为 baseline 处理。")

    add_heading(doc, "Round 2：蓝方 best response", 2)
    add_para(doc, "冻结初始红方，只开放蓝方参数。方法上重点调节 blue_desired_speed、blue_risk_margin、blue_lane_spacing 和蓝方防守机 intercept 行为。目标是让蓝方赛车保持穿门效率，同时由防守机压制红方赛车。结果 blue-red=4.000，蓝方达标。")

    add_heading(doc, "Round 3：红方 best response", 2)
    add_para(doc, "冻结 Round 2 蓝方，只开放红方参数并加入红方 inter-team buffer。方法上提升红方推进和受压避让能力，避免蓝方拦截导致红方减速或碰撞。结果 red-blue=0.667，红方小幅达标。")

    add_heading(doc, "Round 4：蓝方继续改进", 2)
    add_para(doc, "冻结 Round 3 红方。蓝方先尝试 intercept pressure layer，让蓝方防守机压迫领先红方赛车；在初始 3 次代码迭代内曾失败。代码迭代上限提高到 10 后，继续执行，新增一次蓝方 gate-frame guard，并采用更稳定的 escort/guard 组合。结果 blue-red=3.333，蓝方达标。")

    add_heading(doc, "Round 5：红方继续改进", 2)
    add_para(doc, "冻结 Round 4 蓝方 best trial。红方尝试 red breakout drive、escape gain、defender screen、vertical split-lane、gate-frame guard 和 comeback boost。红方效用从 5.0 提升到 9.0，但最优结果仍为 red-blue=-0.333，未能反超蓝方。停止原因是 10/10 次红方受控策略代码迭代已用尽。")

    add_heading(doc, "5. 可视化结果", 1)
    add_picture_if_exists(
        doc,
        FIG_DIR / "current_round_dashboard.png",
        "图 1：Round 2-5 总览。Round 5 的目标 margin 为负，说明红方未反超；硬约束全程为 0。",
        width=6.6,
    )
    add_picture_if_exists(
        doc,
        FIG_DIR / "trial_advantage_distribution.png",
        "图 2：trial 级 advantage 分布。Round 5 虽有多个安全可行 trial，但 promotion gate 未通过。",
        width=6.3,
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "6. 3D 轨迹示例", 1)
    add_picture_if_exists(
        doc,
        FIG_DIR / "trajectory_3d_round02_blue_pass_seed23.png",
        "图 3：Round 2 蓝方 PASS 的 seed 23 轨迹。蓝方在初始红方冻结条件下形成明显优势。",
        width=6.3,
    )
    add_picture_if_exists(
        doc,
        FIG_DIR / "trajectory_3d_round03_red_pass_seed23.png",
        "图 4：Round 3 红方 PASS 的 seed 23 轨迹。红方通过推进与避让改进实现小幅反超。",
        width=6.3,
    )
    add_picture_if_exists(
        doc,
        FIG_DIR / "trajectory_3d_round04_blue_pass_seed23.png",
        "图 5：Round 4 蓝方 PASS 后的 seed 23 轨迹。蓝方 escort/guard 策略取得明显优势。",
        width=6.3,
    )
    add_picture_if_exists(
        doc,
        FIG_DIR / "trajectory_3d_round05_red_fail_stop_seed23.png",
        "图 6：Round 5 红方 FAIL_STOP 的 seed 23 轨迹。该 seed 红方表现较好，但跨 seed 平均仍未反超。",
        width=6.3,
    )

    add_heading(doc, "7. 科研阐述", 1)
    add_para(doc, "科研结论一：该任务存在有效对抗耦合。各轮目标侧空场耦合 Delta 均大于 0，说明对手存在会降低目标侧效用，任务不是两个独立单队竞速问题。")
    add_para(doc, "科研结论二：在当前规则策略空间和 vertical_wave_3v3 布局下，蓝方策略族更稳。Round 5 中红方经过 10 次代码迭代仍未达到 red_utility - blue_utility > 0。")
    add_para(doc, "科研结论三：红方存在显著 best-response 改进空间。Round 5 的 best_response_gain_red=4.000，说明红方策略改进有效，但不足以跨 seed 稳定反超。")
    add_para(doc, "科研结论四：当前策略组合不能声明经验近似纳什稳定。最近 best_response_gain_red=4.000，best_response_gain_blue=0.333，至少一方仍有正收益改进空间。")
    add_para(doc, "推荐表述：在 vertical_wave_3v3 对抗竞速任务中，基于安全规则控制器的交替最佳响应显示，蓝方 escort/guard 型策略在有限搜索空间内对红方 10 次受控最佳响应保持轻微优势；但红方仍存在显著 best-response gain，因此当前策略组合不满足经验近似纳什稳定。")

    add_heading(doc, "8. 失败原因与后续方向", 1)
    add_para(doc, "Round 5 失败的直接原因是最优 trial 的平均效用仍为 red_utility=9.000、blue_utility=9.333，即 red-blue=-0.333。硬约束和 Delta_R 均通过，因此失败不是安全性问题，而是目标优势未达成。")
    add_bullet(doc, "关键薄弱点：seed 11 中红方 6、蓝方 14，红方大幅落后，拉低跨 seed 平均。")
    add_bullet(doc, "可能后续方向：针对 seed 11 的出生位置和早期 gate 竞争做定向诊断。")
    add_bullet(doc, "方法方向：引入更系统的对手建模或短时规划，而不是继续堆叠几何规则 patch。")
    add_bullet(doc, "实验方向：扩大 seeds、记录 per-agent gate sequence、near miss 和 defender-screen 事件，以定位红方失败模式。")

    add_heading(doc, "9. 关键产物与验证", 1)
    add_bullet(doc, "Round history：game/vertical_wave_3v3/round_history.json")
    add_bullet(doc, "最终报告：game/vertical_wave_3v3/summary.md")
    add_bullet(doc, "Round 4 experiment：experiments/vertical_wave_3v3_exp_001_r04_blue/")
    add_bullet(doc, "Round 5 experiment：experiments/vertical_wave_3v3_exp_001_r05_red/")
    add_bullet(doc, "可视化目录：game/vertical_wave_3v3/figures/")
    add_para(doc, "最终验证命令已通过：Round 4 与 Round 5 的 policy hook、pytest 和 post_experiment_run 均通过。")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_doc()
